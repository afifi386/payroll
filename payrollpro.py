import sys
import os
import sqlite3
from datetime import datetime
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, 
                             QLabel, QLineEdit, QComboBox, QPushButton, QTableWidget, 
                             QTableWidgetItem, QScrollArea, QFrame, QFileDialog, QMessageBox,
                             QTabWidget, QGridLayout, QGroupBox, QHeaderView, QMenuBar, QMenu,
                             QAction, QDialog, QFormLayout, QSpinBox, QDoubleSpinBox)
from PyQt5.QtGui import QFont, QIcon, QPixmap
from PyQt5.QtCore import Qt, QSize
from PyQt5.QtPrintSupport import QPrinter, QPrintDialog, QPrintPreviewDialog
import pandas as pd
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.colors import grey, whitesmoke, beige, black
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import arabic_reshaper
from bidi.algorithm import get_display

class PayrollApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("نظام إدارة الرواتب")
        self.setGeometry(100, 100, 1000, 800)
        self.setLayoutDirection(Qt.RightToLeft)  # Set layout direction to RTL
        
        # Create font for Arabic
        self.arabic_font = QFont("Arial", 10)
        self.setFont(self.arabic_font)
        
        # Initialize database
        self.init_database()
        
        # Create UI
        self.create_menu_bar()
        self.create_main_ui()
        
        # Initialize results
        self.results = {}
        self.raw_results = None
        
    def create_menu_bar(self):
        # Create menu bar
        menubar = self.menuBar()
        menubar.setFont(QFont("Arial", 10))
        
        # File menu
        file_menu = menubar.addMenu("ملف")
        
        # New action
        new_action = QAction("جديد", self)
        new_action.setShortcut("Ctrl+N")
        new_action.triggered.connect(self.new_record)
        file_menu.addAction(new_action)
        
        # Save action
        save_action = QAction("حفظ إلى قاعدة البيانات", self)
        save_action.setShortcut("Ctrl+S")
        save_action.triggered.connect(self.save_to_database)
        file_menu.addAction(save_action)
        
        file_menu.addSeparator()
        
        # Export submenu
        export_menu = file_menu.addMenu("تصدير")
        
        export_pdf_action = QAction("تصدير PDF", self)
        export_pdf_action.triggered.connect(self.export_to_pdf)
        export_menu.addAction(export_pdf_action)
        
        export_excel_action = QAction("تصدير Excel", self)
        export_excel_action.triggered.connect(self.export_to_excel)
        export_menu.addAction(export_excel_action)
        
        export_word_action = QAction("تصدير Word", self)
        export_word_action.triggered.connect(self.export_to_word)
        export_menu.addAction(export_word_action)
        
        file_menu.addSeparator()
        
        # Exit action
        exit_action = QAction("خروج", self)
        exit_action.setShortcut("Ctrl+Q")
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)
        
        # View menu
        view_menu = menubar.addMenu("عرض")
        
        view_data_action = QAction("عرض بيانات الرواتب", self)
        view_data_action.triggered.connect(self.view_data_from_database)
        view_menu.addAction(view_data_action)
        
        # Settings menu
        settings_menu = menubar.addMenu("إعدادات")
        
        # Font size action
        font_action = QAction("حجم الخط", self)
        font_action.triggered.connect(self.change_font_size)
        settings_menu.addAction(font_action)
        
        # Help menu
        help_menu = menubar.addMenu("مساعدة")
        
        about_action = QAction("حول البرنامج", self)
        about_action.triggered.connect(self.show_about)
        help_menu.addAction(about_action)
        
    def create_main_ui(self):
        # Create central widget
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        # Main layout
        main_layout = QVBoxLayout(central_widget)
        
        # Title label
        title_label = QLabel("نظام إدارة الرواتب")
        title_label.setFont(QFont("Arial", 24, QFont.Bold))
        title_label.setAlignment(Qt.AlignCenter)
        main_layout.addWidget(title_label)
        
        # Create employee details group
        employee_group = QGroupBox("بيانات الموظف")
        employee_group.setFont(QFont("Arial", 12))
        main_layout.addWidget(employee_group)
        
        employee_layout = QGridLayout(employee_group)
        
        # Define common departments and job titles for comboboxes
        self.departments = ["جرافيك", "تصوير", "ديكور", "عمارة"]
        self.job_titles = ["أ.د", "أ.م.د", "د", "م.م", "م"]
        
        # Row 0
        employee_layout.addWidget(QLabel("الاسم:"), 0, 0)
        self.name_entry = QLineEdit()
        self.name_entry.setFont(QFont("Arial", 12))
        self.name_entry.setAlignment(Qt.AlignCenter)
        employee_layout.addWidget(self.name_entry, 0, 1)
        
        employee_layout.addWidget(QLabel("رقم الموظف:"), 0, 2)
        self.id_entry = QLineEdit()
        self.id_entry.setFont(QFont("Arial", 12))
        self.id_entry.setAlignment(Qt.AlignCenter)
        employee_layout.addWidget(self.id_entry, 0, 3)
        
        # Row 1
        employee_layout.addWidget(QLabel("القسم:"), 1, 0)
        self.department_combo = QComboBox()
        self.department_combo.setFont(QFont("Arial", 12))
        self.department_combo.addItems(self.departments)
        employee_layout.addWidget(self.department_combo, 1, 1)
        
        employee_layout.addWidget(QLabel("الدرجة الوظيفية:"), 1, 2)
        self.job_title_combo = QComboBox()
        self.job_title_combo.setFont(QFont("Arial", 12))
        self.job_title_combo.addItems(self.job_titles)
        employee_layout.addWidget(self.job_title_combo, 1, 3)
        
        # Row 2
        employee_layout.addWidget(QLabel("الراتب الأساسي:"), 2, 0)
        self.basic_salary_entry = QLineEdit()
        self.basic_salary_entry.setFont(QFont("Arial", 12))
        self.basic_salary_entry.setAlignment(Qt.AlignCenter)
        employee_layout.addWidget(self.basic_salary_entry, 2, 1)
        
        employee_layout.addWidget(QLabel("اجتماعية:"), 2, 2)
        self.social_entry = QLineEdit()
        self.social_entry.setFont(QFont("Arial", 12))
        self.social_entry.setAlignment(Qt.AlignCenter)
        employee_layout.addWidget(self.social_entry, 2, 3)
        
        # Row 3
        employee_layout.addWidget(QLabel("أساسى 30/6/15:"), 3, 0)
        self.basic30_entry = QLineEdit()
        self.basic30_entry.setFont(QFont("Arial", 12))
        self.basic30_entry.setAlignment(Qt.AlignCenter)
        employee_layout.addWidget(self.basic30_entry, 3, 1)
        
        employee_layout.addWidget(QLabel("إعانة:"), 3, 2)
        self.enaa = QLineEdit()
        self.enaa.setFont(QFont("Arial", 12))
        self.enaa.setAlignment(Qt.AlignCenter)
        employee_layout.addWidget(self.enaa, 3, 3)
        
        # Add button layout
        button_layout = QHBoxLayout()
        main_layout.addLayout(button_layout)
        
        # Create buttons
        self.calculate_button = QPushButton("حساب")
        self.calculate_button.setFont(QFont("Arial", 12))
        self.calculate_button.clicked.connect(self.calculate_payroll)
        button_layout.addWidget(self.calculate_button)
        
        self.save_button = QPushButton("حفظ إلى قاعدة البيانات")
        self.save_button.setFont(QFont("Arial", 12))
        self.save_button.clicked.connect(self.save_to_database)
        button_layout.addWidget(self.save_button)
        
        self.export_pdf_button = QPushButton("تصدير PDF")
        self.export_pdf_button.setFont(QFont("Arial", 12))
        self.export_pdf_button.clicked.connect(self.export_to_pdf)
        button_layout.addWidget(self.export_pdf_button)
        
        self.export_excel_button = QPushButton("تصدير Excel")
        self.export_excel_button.setFont(QFont("Arial", 12))
        self.export_excel_button.clicked.connect(self.export_to_excel)
        button_layout.addWidget(self.export_excel_button)
        
        self.export_word_button = QPushButton("تصدير Word")
        self.export_word_button.setFont(QFont("Arial", 12))
        self.export_word_button.clicked.connect(self.export_to_word)
        button_layout.addWidget(self.export_word_button)
        
        # Create view data button
        view_data_layout = QHBoxLayout()
        main_layout.addLayout(view_data_layout)
        
        self.view_data_button = QPushButton("عرض بيانات الرواتب")
        self.view_data_button.setFont(QFont("Arial", 12))
        self.view_data_button.clicked.connect(self.view_data_from_database)
        view_data_button_container = QWidget()
        view_data_button_container_layout = QHBoxLayout(view_data_button_container)
        view_data_button_container_layout.addWidget(self.view_data_button)
        view_data_layout.addWidget(view_data_button_container)
        
        # Create results group
        results_group = QGroupBox("نتائج الحساب")
        results_group.setFont(QFont("Arial", 12))
        main_layout.addWidget(results_group)
        
        results_layout = QVBoxLayout(results_group)
        
        # Create table for results
        self.results_table = QTableWidget(0, 2)
        self.results_table.setFont(QFont("Arial", 12))
        self.results_table.setHorizontalHeaderLabels(["البند", "القيمة"])
        self.results_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.results_table.verticalHeader().setVisible(False)
        results_layout.addWidget(self.results_table)
        
    def init_database(self):
        conn = sqlite3.connect('payroll.db')
        cursor = conn.cursor()
        cursor.execute('''
        CREATE TABLE IF NOT EXISTS payroll (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            employee_id TEXT,
            employee_name TEXT,
            department TEXT,
            job_title TEXT,
            basic_salary REAL,
            Social REAL,
            basic30 REAL,
            enaa REAL,
            bhos REAL,
            ryada REAL,
            eshraf REAL,
            maktabia REAL,
            tatwer REAL,
            gawda REAL,
            diff_gawda REAL,
            hafz REAL,
            badl REAL,
            salary REAL,
            date TEXT
        )
        ''')
        conn.commit()
        conn.close()
        
    def new_record(self):
        # Clear all input fields
        self.name_entry.clear()
        self.id_entry.clear()
        self.basic_salary_entry.clear()
        self.social_entry.clear()
        self.basic30_entry.clear()
        self.enaa.clear()
        self.department_combo.setCurrentIndex(0)
        self.job_title_combo.setCurrentIndex(0)
        
        # Clear results table
        self.results_table.setRowCount(0)
        
        # Reset results
        self.results = {}
        self.raw_results = None
        
    def calculate_payroll(self):
        try:
            result = 0
            # Get values from entries and comboboxes
            employee_name = self.name_entry.text()
            employee_id = self.id_entry.text()
            department = self.department_combo.currentText()
            job_title = self.job_title_combo.currentText()
            
            # Validate required fields
            if not employee_name or not employee_id:
                QMessageBox.critical(self, "خطأ", "اسم الموظف ورقم الموظف مطلوبان.")
                return
                
            if not department or not job_title:
                QMessageBox.critical(self, "خطأ", "القسم والدرجة الوظيفية مطلوبان.")
                return
            
            # Get and validate numeric inputs
            try:
                basic_salary = float(self.basic_salary_entry.text())
                social = float(self.social_entry.text())
                basic30 = float(self.basic30_entry.text())
                enaa = float(self.enaa.text())
            except ValueError:
                QMessageBox.critical(self, "خطأ", "يرجى إدخال قيم رقمية صحيحة للراتب والساعات ومعدل الضريبة.")
                return
                
            # Calculate payroll
            result += (basic30 * 0.775)
            bhos = (basic30 * 0.49)
            result += bhos
            ryada = (basic30 * 0.91)
            result += ryada
            eshraf = (basic30 * 1.3)
            result += eshraf
            maktabia = (basic30 * 0.78)
            result += maktabia
            tatwer = (basic30 * 0.78)
            result += tatwer
            gawda = 0
            diff_gawda = 0
            hafz = 0
            badl = 0
            exp = 600
            result += exp
            mnha = 10
            result += mnha
            alawa = 73.90
            result += alawa
            tdress = 1071
            result += tdress
            
            # Apply job title specific calculations
            if job_title == "أ.د":
                badl = 3500
                result += badl
                gawda = 4270
                diff_gawda = 330
                hafz = 2600
                result += (gawda + diff_gawda + hafz)
            elif job_title == "أ.م.د":
                badl = 3000
                result += badl
                gawda = 3770
                diff_gawda = 230
                hafz = 2475
                result += (gawda + diff_gawda + hafz)
            elif job_title == "د":
                badl = 2500
                result += badl
                gawda = 3120
                diff_gawda = 140
                hafz = 2050
                result += (gawda + diff_gawda + hafz)
            elif job_title == "م.م":
                badl = 1500
                result += badl
                gawda = 2900
                diff_gawda = 100
                hafz = 1850
                result += (gawda + diff_gawda + hafz)
            elif job_title == "م":
                badl = 1000
                result += badl
                gawda = 1850
                diff_gawda = 40
                hafz = 1850
                result += (gawda + diff_gawda + hafz)
                
            # Store results in dictionary
            self.results = {
                "رقم الموظف": employee_id,
                "الاسم": employee_name,
                "القسم": department,
                "الدرجة الوظيفية": job_title,
                "الاساسى": f"${basic_salary:.2f}",
                "اجتماعية": f"{social}",
                "اساسى 30/6/15": f"{basic30}",
                "اعانة": f"{enaa}",
                "بحوث": f"${bhos:.2f}",
                "ريادة": f"${ryada:.2f}",
                "اشراف": f"${eshraf:.2f}",
                "مكتبية": f"${maktabia:.2f}",
                "تطوير": f"${tatwer:.2f}",
                "جودة": f"${gawda:.2f}",
                "فرق الجودة": f"${diff_gawda:.2f}",
                "حافز": f"${hafz:.2f}",
                "بدل": f"${badl:.2f}",
                "جملة الاجر": f"${result:.2f}"
            }

            # Save raw values for database
            self.raw_results = {
                "employee_id": employee_id,
                "employee_name": employee_name,
                "department": department,
                "job_title": job_title,
                "basic_salary": basic_salary,
                "Social": social,
                "basic30": basic30,
                "enaa": enaa,
                "bhos": bhos,
                "ryada": ryada,
                "eshraf": eshraf,
                "maktabia": maktabia,
                "tatwer": tatwer,
                "gawda": gawda,
                "diff_gawda": diff_gawda,
                "hafz": hafz,
                "badl": badl,
                "salary": result
            }

            # Display results
            self.display_results()
            
        except Exception as e:
            QMessageBox.critical(self, "خطأ", f"حدث خطأ أثناء الحساب: {str(e)}")
    
    def display_results(self):
        # Clear previous results
        self.results_table.setRowCount(0)
        
        # Add results to table
        for key, value in self.results.items():
            row_position = self.results_table.rowCount()
            self.results_table.insertRow(row_position)
            
            # Set key item
            key_item = QTableWidgetItem(key)
            key_item.setFont(QFont("Arial", 11, QFont.Bold))
            self.results_table.setItem(row_position, 0, key_item)
            
            # Set value item
            value_item = QTableWidgetItem(value)
            value_item.setFont(QFont("Arial", 11))
            self.results_table.setItem(row_position, 1, value_item)
    
    def save_to_database(self):
        if not self.raw_results:
            QMessageBox.critical(self, "خطأ", "قم بحساب الراتب أولاً قبل الحفظ في قاعدة البيانات.")
            return
        
        try:
            conn = sqlite3.connect('payroll.db')
            cursor = conn.cursor()
            
            # Current date
            current_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            
            # Insert data
            cursor.execute('''
            INSERT INTO payroll (
                employee_id, employee_name, department, job_title, basic_salary, Social, 
                basic30, enaa, bhos, ryada, 
                eshraf, maktabia, tatwer, gawda, diff_gawda, hafz, badl, salary, date
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                self.raw_results["employee_id"],
                self.raw_results["employee_name"],
                self.raw_results["department"],
                self.raw_results["job_title"],
                self.raw_results["basic_salary"],
                self.raw_results["Social"],
                self.raw_results["basic30"],
                self.raw_results["enaa"],
                self.raw_results["bhos"],
                self.raw_results["ryada"],
                self.raw_results["eshraf"],
                self.raw_results["maktabia"],
                self.raw_results["tatwer"],
                self.raw_results["gawda"],
                self.raw_results["diff_gawda"],
                self.raw_results["hafz"],
                self.raw_results["badl"],
                self.raw_results["salary"],
                current_date
            ))
            
            conn.commit()
            conn.close()
            
            QMessageBox.information(self, "نجاح", "تم حفظ بيانات الراتب في قاعدة البيانات بنجاح.")
        except Exception as e:
            QMessageBox.critical(self, "خطأ قاعدة البيانات", f"فشل الحفظ في قاعدة البيانات: {e}")
    
    def view_data_from_database(self):
        # Create a new dialog window
        data_dialog = QDialog(self)
        data_dialog.setWindowTitle("بيانات الرواتب")
        data_dialog.setGeometry(100, 100, 1200, 700)
        data_dialog.setLayoutDirection(Qt.RightToLeft)
        
        # Main layout
        main_layout = QVBoxLayout(data_dialog)
        
        # Create search frame
        search_frame = QFrame()
        search_layout = QHBoxLayout(search_frame)
        
        search_layout.addWidget(QLabel("البحث حسب:"))
        
        search_options = ["رقم الموظف", "اسم الموظف", "القسم", "الدرجة الوظيفية"]
        self.search_by = QComboBox()
        self.search_by.addItems(search_options)
        search_layout.addWidget(self.search_by)
        
        search_layout.addWidget(QLabel("مصطلح البحث:"))
        self.search_entry = QLineEdit()
        search_layout.addWidget(self.search_entry)
        
        self.search_button = QPushButton("بحث")
        self.search_button.clicked.connect(lambda: self.search_data(data_table))
        search_layout.addWidget(self.search_button)
        
        self.reset_button = QPushButton("إعادة تعيين")
        self.reset_button.clicked.connect(lambda: self.load_all_data(data_table))
        search_layout.addWidget(self.reset_button)
        
        main_layout.addWidget(search_frame)
        
        # Create table widget
        data_table = QTableWidget()
        data_table.setFont(QFont("Arial", 10))
        
        # Set column headers
        columns = ["رقم الموظف", "الاسم", "القسم", "الدرجة", "الاساسى", 
                   "اجتماعية", "اساسى30/6/15", "اعانه", "بحوث", "ريادة", "اشراف", "مكتبية", 
                   "تطوير", "جودة", "فرق الجودة", "حافز", "بدل", "جملة الاجر", "التاريخ"]
        data_table.setColumnCount(len(columns))
        data_table.setHorizontalHeaderLabels(columns)
        data_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)
        
        main_layout.addWidget(data_table)
        
        # Create button frame
        button_frame = QFrame()
        button_layout = QHBoxLayout(button_frame)
        
        export_excel_btn = QPushButton("تصدير إلى Excel")
        export_excel_btn.clicked.connect(lambda: self.export_view_to_excel(data_table))
        button_layout.addWidget(export_excel_btn)
        
        export_pdf_btn = QPushButton("تصدير إلى PDF")
        export_pdf_btn.clicked.connect(lambda: self.export_view_to_pdf(data_table))
        button_layout.addWidget(export_pdf_btn)
        
        export_word_btn = QPushButton("تصدير إلى Word")
        export_word_btn.clicked.connect(lambda: self.export_view_to_word(data_table))
        button_layout.addWidget(export_word_btn)
        
        main_layout.addWidget(button_frame)
        
        # Load initial data
        self.load_all_data(data_table)
        
        # Show dialog
        data_dialog.exec_()
    
    def load_all_data(self, table):
        # Clear the table
        table.setRowCount(0)
        
        # Fetch all data from the database
        conn = sqlite3.connect('payroll.db')
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM payroll")
        rows = cursor.fetchall()
        conn.close()
        
        # Insert rows into the table
        for row in rows:
            row_position = table.rowCount()
            table.insertRow(row_position)
            
            # Skip the first column (ID)
            for i, value in enumerate(row[1:]):
                item = QTableWidgetItem(str(value))
                table.setItem(row_position, i, item)
    
    def search_data(self, table):
        search_column = self.search_by.currentText()
        search_term = self.search_entry.text()
        
        if not search_term:
            self.load_all_data(table)
            return
        
        # Map UI column names to database column names
        column_mapping = {
            "رقم الموظف": "employee_id",
            "اسم الموظف": "employee_name",
            "القسم": "department",
            "الدرجة الوظيفية": "job_title"
        }
        
        # Get the database column name
        db_column = column_mapping.get(search_column)
        
        if not db_column:
            QMessageBox.critical(self, "خطأ", "تم تحديد عمود بحث غير صالح.")
            return
        
        # Clear the table
        table.setRowCount(0)
        
        # Search in the database
        conn = sqlite3.connect('payroll.db')
        cursor = conn.cursor()
        cursor.execute(f"SELECT * FROM payroll WHERE {db_column} LIKE ?", (f"%{search_term}%",))
        rows = cursor.fetchall()
        conn.close()
        
        # Insert matching rows into the table
        for row in rows:
            row_position = table.rowCount()
            table.insertRow(row_position)
            
            # Skip the first column (ID)
            for i, value in enumerate(row[1:]):
                item = QTableWidgetItem(str(value))
                table.setItem(row_position, i, item)
    
    def export_view_to_excel(self, table):
        # Get selected rows
        selected_items = table.selectedItems()
        if not selected_items:
            QMessageBox.information(self, "معلومات", "الرجاء تحديد صفوف للتصدير.")
            return
        
        try:
            # Ask user where to save the file
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "حفظ تقرير Excel",
                "",
                "Excel files (*.xlsx)"
            )
            
            if not file_path:
                return  # User cancelled
                
            # Get column headers
            headers = []
            for i in range(table.columnCount()):
                headers.append(table.horizontalHeaderItem(i).text())
            
            # Get selected rows (without duplicates)
            selected_rows = sorted(set(item.row() for item in selected_items))
            
            # Get data from selected rows
            data = []
            for row in selected_rows:
                row_data = []
                for col in range(table.columnCount()):
                    item = table.item(row, col)
                    if item is not None:
                        row_data.append(item.text())
                    else:
                        row_data.append("")
                data.append(row_data)
            
            # Create DataFrame and save to Excel
            df = pd.DataFrame(data, columns=headers)
            
            # Create Excel writer
            with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
                # Write DataFrame to Excel
                df.to_excel(writer, index=False, sheet_name="بيانات الرواتب")
                
                # Get the workbook and worksheet
                workbook = writer.book
                worksheet = writer.sheets["بيانات الرواتب"]
                
                # Adjust column widths for better readability
                for i, col in enumerate(df.columns):
                    column_width = max(df[col].astype(str).map(len).max(), len(col)) + 2
                    col_letter = chr(65 + i) if i < 26 else chr(64 + i // 26) + chr(65 + i % 26)
                    worksheet.column_dimensions[col_letter].width = column_width
            
            QMessageBox.information(self, "نجاح", f"تم تصدير البيانات المحددة بنجاح إلى {file_path}")
            
        except Exception as e:
            QMessageBox.critical(self, "خطأ", f"حدث خطأ أثناء التصدير إلى Excel: {str(e)}")
    
    def export_view_to_pdf(self, table):
        # Get selected rows
        selected_items = table.selectedItems()
        if not selected_items:
            QMessageBox.information(self, "معلومات", "الرجاء تحديد صفوف للتصدير.")
            return
        
        try:
            # Ask user where to save the file
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "حفظ تقرير PDF",
                "",
                "PDF files (*.pdf)"
            )
            
            if not file_path:
                return  # User cancelled
            
            # Register the Arabic font
            pdfmetrics.registerFont(TTFont('Arabic', 'Amiri-Regular.ttf'))
            
            # Create a PDF document
            doc = SimpleDocTemplate(file_path, pagesize=letter)
            
            # Define styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'Title',
                parent=styles['Heading1'],
                fontName='Arabic',
                fontSize=18,
                alignment=1,  # Center alignment
                spaceAfter=12
            )
            
            subtitle_style = ParagraphStyle(
                'Subtitle',
                parent=styles['Heading2'],
                fontName='Arabic',
                fontSize=14,
                alignment=1,  # Center alignment
                spaceAfter=10
            )
            
            # Create document elements
            elements = []
            
            # Add main title with proper Arabic text handling
            title_text = arabic_reshaper.reshape("تقرير الرواتب")
            title_bidi = get_display(title_text)
            title = Paragraph(title_bidi, title_style)
            elements.append(title)
            elements.append(Spacer(1, 0.5 * inch))  # Increased spacing after main title
            
            # Get headers
            headers = []
            for col in range(table.columnCount()):
                text = table.horizontalHeaderItem(col).text()
                # Use Arabic reshaper to fix text display
                reshaped_text = arabic_reshaper.reshape(text)
                bidi_text = get_display(reshaped_text)
                headers.append(bidi_text)
            
            # Find the name column index (assuming there's a column with الاسم or اسم الموظف)
            name_column_index = 0  # Default to the first column
            for i, header in enumerate(headers):
                # Check for Arabic words for "name" in their display form
                if "اسم" in get_display(arabic_reshaper.reshape("اسم")) in header or "الاسم" in header:
                    name_column_index = i
                    break
            
            # Get selected rows (without duplicates)
            selected_rows = sorted(set(item.row() for item in selected_items))
            
            # Process each selected row vertically
            for row in selected_rows:
                # Get employee name from the appropriate column
                name_item = table.item(row, name_column_index)
                
                if name_item is not None:
                    employee_name = name_item.text()
                else:
                    employee_name = f"موظف {row+1}"  # Fallback if name not found
                
                # Create subtitle for this employee record with proper Arabic handling
                subtitle_text = arabic_reshaper.reshape(f"بيانات الموظف: {employee_name}")
                
                subtitle_bidi = get_display(subtitle_text)
                employee_title = Paragraph(subtitle_bidi, subtitle_style)
                elements.append(employee_title)
                elements.append(Spacer(1, 0.25 * inch))
                
                # Create data for vertical display - each header paired with its value
                vertical_data = []
                for col in range(table.columnCount()):
                    header = headers[col]
                    
                    # Get the value
                    item = table.item(row, col)
                    value = ""
                    if item is not None:
                        value_text = item.text()
                        # Use Arabic reshaper to fix text display
                        reshaped_value = arabic_reshaper.reshape(value_text)
                        value = get_display(reshaped_value)
                    
                    # Add header-value pair
                    vertical_data.append([header, value])
                
                # Create vertical table for this employee
                vertical_table = Table(vertical_data)
                
                # Style the vertical table
                v_style = TableStyle([
                    ('BACKGROUND', (0, 0), (0, -1), grey),
                    ('TEXTCOLOR', (0, 0), (0, -1), (1, 1, 1)),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, -1), 'Arabic'),  # Changed to Arabic font
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('BACKGROUND', (1, 0), (1, -1), beige),
                    ('GRID', (0, 0), (-1, -1), 1, black)
                ])
                
                # Add alternating row colors
                for i in range(len(vertical_data)):
                    if i % 2 == 0:
                        v_style.add('BACKGROUND', (1, i), (1, i), whitesmoke)
                
                vertical_table.setStyle(v_style)
                elements.append(vertical_table)
                elements.append(Spacer(1, 0.5 * inch))  # Increased spacing between employee records
            
            # Build PDF
            doc.build(elements)
            
            QMessageBox.information(self, "نجاح", f"تم تصدير البيانات المحددة بنجاح إلى {file_path}")
            
        except Exception as e:
            QMessageBox.critical(self, "خطأ", f"حدث خطأ أثناء التصدير إلى PDF: {str(e)}")
    
    def export_view_to_word(self, table):
        # Get selected rows
        selected_items = table.selectedItems()
        if not selected_items:
            QMessageBox.information(self, "معلومات", "الرجاء تحديد صفوف للتصدير.")
            return
        
        try:
            # Ask user where to save the file
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "حفظ تقرير Word",
                "",
                "Word files (*.docx)"
            )
            
            if not file_path:
                return  # User cancelled
            
            # Import necessary components
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Inches, Pt
            
            # Create a new Word document
            doc = Document()
            
            # Add a title
            title = doc.add_heading("تقرير الرواتب", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Get headers
            headers = []
            for col in range(table.columnCount()):
                header_text = table.horizontalHeaderItem(col).text()
                headers.append(header_text)
            
            # Find the name column index (assuming there's a column with الاسم or اسم الموظف)
            name_column_index = 0  # Default to the first column
            for i, header in enumerate(headers):
                if "اسم" in header or "الاسم" in header:
                    name_column_index = i
                    break
            
            # Get selected rows (without duplicates)
            selected_rows = sorted(set(item.row() for item in selected_items))
            
            # Process each selected employee/row vertically
            for row in selected_rows:
                # Get employee name from the appropriate column
                name_item = table.item(row, name_column_index)
                employee_name = name_item.text() if name_item is not None else f"موظف {row+1}"
                
                # Add a subtitle for this employee
                subtitle = doc.add_heading(f"بيانات الموظف: {employee_name}", level=1)
                subtitle.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                # Create a table with 2 columns (header and value) and rows equal to number of fields
                vertical_table = doc.add_table(rows=table.columnCount(), cols=2)
                vertical_table.style = 'Table Grid'
                
                # Add data in vertical format
                for col in range(table.columnCount()):
                    # Set header in first column
                    header_cell = vertical_table.cell(col, 0)
                    header_cell.text = headers[col]
                    
                    # Format the header cell - make it bold and add background color
                    for paragraph in header_cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.rtl = True
                    
                    # Set value in second column
                    value_cell = vertical_table.cell(col, 1)
                    item = table.item(row, col)
                    if item is not None:
                        value_cell.text = item.text()
                    else:
                        value_cell.text = ""
                    
                    # Format the value cell
                    for paragraph in value_cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        for run in paragraph.runs:
                            run.font.rtl = True
                
                # Add space after each employee's data
                doc.add_paragraph()
            
            # Set document RTL properties globally
            for paragraph in doc.paragraphs:
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in paragraph.runs:
                    run.font.rtl = True
            
            # Save the document
            doc.save(file_path)
            
            QMessageBox.information(self, "نجاح", f"تم تصدير البيانات المحددة بنجاح إلى {file_path}")
            
        except Exception as e:
            QMessageBox.critical(self, "خطأ", f"حدث خطأ أثناء التصدير إلى Word: {str(e)}")
    
    def export_to_pdf(self):
        if not self.results:
            QMessageBox.critical(self, "خطأ", "قم بحساب الراتب أولاً قبل التصدير.")
            return
        
        try:
            # Ask user where to save the file
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "حفظ مرتب كملف PDF",
                "",
                "PDF files (*.pdf)"
            )
            
            if not file_path:
                return  # User cancelled
            
            # Register the Arabic font
            pdfmetrics.registerFont(TTFont('Arabic', 'Amiri-Regular.ttf'))
            
            # Create a PDF document
            doc = SimpleDocTemplate(file_path, pagesize=letter)
            
            # Define styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'Title',
                parent=styles['Heading1'],
                fontName='Arial',
                fontSize=18,
                alignment=1,  # Center alignment
                spaceAfter=12
            )
            
            # Create document elements
            elements = []
            
            # Add title with proper Arabic text handling
            title_text = arabic_reshaper.reshape("بيان مرتب")
            title_bidi = get_display(title_text)
            title = Paragraph(title_bidi, title_style)
            elements.append(title)
            elements.append(Spacer(1, 0.25 * inch))
            
            # Convert results to data for table
            data = []
            for key, value in self.results.items():
                # Use Arabic reshaper to fix text display
                reshaped_key = arabic_reshaper.reshape(key)
                bidi_key = get_display(reshaped_key)
                
                if isinstance(value, str):
                    reshaped_value = arabic_reshaper.reshape(value)
                    bidi_value = get_display(reshaped_value)
                else:
                    bidi_value = str(value)
                
                data.append([bidi_key, bidi_value])
            
            # Create table
            table = Table(data, colWidths=[3*inch, 2*inch])
            
            # Style the table
            style = TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), grey),
                ('TEXTCOLOR', (0, 0), (0, -1), (1, 1, 1)),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, -1), 'Arial'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('BACKGROUND', (1, 0), (1, -1), beige),
                ('GRID', (0, 0), (-1, -1), 1, black)
            ])
            
            # Add alternating row colors
            for i in range(len(data)):
                if i % 2 == 0:
                    style.add('BACKGROUND', (1, i), (1, i), whitesmoke)
            
            table.setStyle(style)
            elements.append(table)
            
            # Add current date with proper Arabic text handling
            current_date = datetime.now().strftime("%Y-%m-%d")
            date_text = arabic_reshaper.reshape(f"تاريخ الإصدار: {current_date}")
            date_bidi = get_display(date_text)
            date_paragraph = Paragraph(date_bidi, styles['Normal'])
            elements.append(Spacer(1, 0.5 * inch))
            elements.append(date_paragraph)
            
            # Build PDF
            doc.build(elements)
            
            QMessageBox.information(self, "نجاح", f"تم تصدير المرتب بنجاح إلى {file_path}")
            
        except Exception as e:
            QMessageBox.critical(self, "خطأ", f"حدث خطأ أثناء التصدير إلى PDF: {str(e)}")
    
    def export_to_excel(self):
        if not self.results:
            QMessageBox.critical(self, "خطأ", "قم بحساب الراتب أولاً قبل التصدير.")
            return
        
        try:
            # Ask user where to save the file
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "حفظ مرتب كملف Excel",
                "",
                "Excel files (*.xlsx)"
            )
            
            if not file_path:
                return  # User cancelled
            
            # Create a single-row DataFrame with items as columns
            # This inverts the data orientation from vertical to horizontal
            items = list(self.results.keys())
            values = list(self.results.values())
            
            # Create a DataFrame with one row, where each column is an item
            df = pd.DataFrame([values], columns=items)
            
            # Create Excel writer
            with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
                # Write DataFrame to Excel
                df.to_excel(writer, index=False, sheet_name="مرتب")
                
                # Get the workbook and worksheet
                workbook = writer.book
                worksheet = writer.sheets["مرتب"]
                
                # Adjust column widths for better readability
                for i, col in enumerate(df.columns):
                    column_width = max(len(str(df[col].iloc[0])), len(col)) + 2
                    # Use Excel column letters (A, B, C, ..., Z, AA, AB, ...)
                    col_letter = chr(65 + i) if i < 26 else chr(65 + i//26 - 1) + chr(65 + i%26)
                    worksheet.column_dimensions[col_letter].width = column_width
            
            QMessageBox.information(self, "نجاح", f"تم تصدير المرتب بنجاح إلى {file_path}")
            
        except Exception as e:
            QMessageBox.critical(self, "خطأ", f"حدث خطأ أثناء التصدير إلى Excel: {str(e)}")
    
    def export_to_word(self):
        if not self.results:
            QMessageBox.critical(self, "خطأ", "قم بحساب الراتب أولاً قبل التصدير.")
            return
        
        try:
            # Ask user where to save the file
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "حفظ مرتب كملف Word",
                "",
                "Word files (*.docx)"
            )
            
            if not file_path:
                return  # User cancelled
            
            # Create a new Word document
            doc = Document()
            
            # Add a title
            doc.add_heading("بيان مرتب", 0)
            
            # Create a table with the appropriate number of rows and columns
            table = doc.add_table(rows=len(self.results) + 1, cols=2)
            table.style = 'Table Grid'
            
            # Add header row
            header_cells = table.rows[0].cells
            header_cells[0].text = "البند"
            header_cells[1].text = "القيمة"
            
            # Add data rows
            for i, (key, value) in enumerate(self.results.items(), 1):
                row_cells = table.rows[i].cells
                row_cells[0].text = key
                row_cells[1].text = str(value)
            
            # Add current date
            doc.add_paragraph("")
            current_date = datetime.now().strftime("%Y-%m-%d")
            doc.add_paragraph(f"تاريخ الإصدار: {current_date}")
            
            # Save the document
            doc.save(file_path)
            
            QMessageBox.information(self, "نجاح", f"تم تصدير المرتب بنجاح إلى {file_path}")
            
        except Exception as e:
            QMessageBox.critical(self, "خطأ", f"حدث خطأ أثناء التصدير إلى Word: {str(e)}")
    
    def change_font_size(self):
        try:
            # Create a dialog for font size selection
            dialog = QDialog(self)
            dialog.setWindowTitle("تغيير حجم الخط")
            dialog.setFixedSize(300, 150)
            
            layout = QVBoxLayout(dialog)
            
            form_layout = QFormLayout()
            
            # Font size spinner
            font_size_spin = QSpinBox()
            font_size_spin.setRange(8, 20)
            font_size_spin.setValue(self.arabic_font.pointSize())
            form_layout.addRow("حجم الخط:", font_size_spin)
            
            layout.addLayout(form_layout)
            
            # Buttons
            button_layout = QHBoxLayout()
            
            ok_button = QPushButton("موافق")
            ok_button.clicked.connect(dialog.accept)
            button_layout.addWidget(ok_button)
            
            cancel_button = QPushButton("إلغاء")
            cancel_button.clicked.connect(dialog.reject)
            button_layout.addWidget(cancel_button)
            
            layout.addLayout(button_layout)
            
            # Show dialog
            result = dialog.exec_()
            
            if result == QDialog.Accepted:
                # Get new font size
                new_size = font_size_spin.value()
                
                # Update font
                self.arabic_font.setPointSize(new_size)
                self.setFont(self.arabic_font)
                
                # Update UI components
                self.update_font_size_in_ui(new_size)
                
                QMessageBox.information(self, "نجاح", "تم تغيير حجم الخط بنجاح.")
                
        except Exception as e:
            QMessageBox.critical(self, "خطأ", f"حدث خطأ أثناء تغيير حجم الخط: {str(e)}")
    
    def update_font_size_in_ui(self, size):
        # Update fonts in all UI elements
        font = QFont("Arial", size)
        
        # Update input fields
        self.name_entry.setFont(font)
        self.id_entry.setFont(font)
        self.basic_salary_entry.setFont(font)
        self.social_entry.setFont(font)
        self.basic30_entry.setFont(font)
        self.enaa.setFont(font)
        
        # Update comboboxes
        self.department_combo.setFont(font)
        self.job_title_combo.setFont(font)
        
        # Update buttons
        self.calculate_button.setFont(font)
        self.save_button.setFont(font)
        self.export_pdf_button.setFont(font)
        self.export_excel_button.setFont(font)
        self.export_word_button.setFont(font)
        self.view_data_button.setFont(font)
        
        # Update results table
        self.results_table.setFont(font)
    
    def show_about(self):
        about_text = """
        <h1>نظام إدارة الرواتب</h1>
        <p>تطبيق لإدارة وحساب رواتب الموظفين</p>
        <p>الإصدار 1.0</p>
        <strong> تطوير مطور محمد حسين</strong>
        <p>جميع الحقوق محفوظة © 2025</p>
        """
        
        QMessageBox.about(self, "حول البرنامج", about_text)

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = PayrollApp()
    window.show()
    sys.exit(app.exec_())