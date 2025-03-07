from tkinter import Canvas, Toplevel, Tk
from tkinter import messagebox, filedialog
from tkinter.ttk import Button, Scrollbar, Style, Frame, Treeview
from tkinter.ttk import Label, LabelFrame, Combobox, Entry
import sqlite3
from datetime import datetime
from reportlab.lib.colors import grey, whitesmoke, beige, black
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from reportlab.lib.units import inch
from pandas import DataFrame, ExcelWriter
from docx import Document
import openpyxl
from reportlab.platypus import Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
import os
from bidi.algorithm import get_display  # For RTL text reordering
import arabic_reshaper  # For reshaping Arabic text

class PayrollApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Payroll Management System")
        self.root.geometry("800x600")
        
        # Set theme and style
        style = Style()
        style.theme_use('clam')  # You can choose: 'clam', 'alt', 'default', 'classic'
        
        # Initialize database
        self.init_database()
        
        # Main frame
        self.main_frame = Frame(self.root, padding=20)
        self.main_frame.pack(fill="both", expand=True)
        
        # Title
        title_label = Label(
            self.main_frame, 
            text="Payroll Management System", 
            font=("Helvetica", 24, "bold")
        )
        title_label.pack(pady=20)
        
        # Create entry frame
        self.entry_frame = LabelFrame(self.main_frame, text="Employee Details", padding=10)
        self.entry_frame.pack(fill="x", padx=10, pady=20)
        
        # Create input fields
        self.create_input_fields()
        
        # Create buttons
        self.button_frame = Frame(self.main_frame)
        self.button_frame.pack(fill="x", padx=10, pady=20)
        
        self.calculate_button = Button(
            self.button_frame,
            text="Calculate",
            command=self.calculate_payroll
        )
        self.calculate_button.grid(row=0, column=0, padx=5, pady=10)
        
        self.save_button = Button(
            self.button_frame,
            text="Save to Database",
            command=self.save_to_database
        )
        self.save_button.grid(row=0, column=1, padx=5, pady=10)
        
        self.export_pdf_button = Button(
            self.button_frame,
            text="Export to PDF",
            command=self.export_to_pdf
        )
        self.export_pdf_button.grid(row=0, column=2, padx=5, pady=10)
        
        self.export_excel_button = Button(
            self.button_frame,
            text="Export to Excel",
            command=self.export_to_excel
        )
        self.export_excel_button.grid(row=0, column=3, padx=5, pady=10)
        
        self.export_word_button = Button(
            self.button_frame,
            text="Export to Word",
            command=self.export_to_word
        )
        self.export_word_button.grid(row=0, column=4, padx=5, pady=10)

        # Button to fetch and display data
        self.view_data_button = Button(
            self.button_frame,
            text="View Payroll Data",
            command=self.view_data_from_database
        )
        self.view_data_button.grid(row=1, column=0, columnspan=5, pady=10)

        # Create display frame for calculations
        self.result_frame = LabelFrame(self.main_frame, text="Calculation Results", padding=10)
        self.result_frame.pack(fill="both", expand=True, padx=10, pady=20)

        # Create canvas with scrollbar for results
        self.canvas = Canvas(self.result_frame)
        self.scrollbar = Scrollbar(self.result_frame, orient="vertical", command=self.canvas.yview)
        self.scrollable_frame = Frame(self.canvas)

        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all"))
        )

        self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        self.canvas.configure(yscrollcommand=self.scrollbar.set)

        self.canvas.pack(side="left", fill="both", expand=True)
        self.scrollbar.pack(side="right", fill="y")

        self.results = {}
        self.result_widgets = {}

    # New function to fetch and display data from the database in another window
    def view_data_from_database(self):
        # Create a new window
        data_window = Toplevel(self.root)
        data_window.title("Payroll Data")
        data_window.geometry("1000x600")
        
        # Create a frame for search controls
        search_frame = Frame(data_window, padding=10)
        search_frame.pack(fill="x", padx=10, pady=10)
        
        # Create search widgets
        Label(search_frame, text="Search by:", font=("Helvetica", 11)).grid(row=0, column=0, padx=5, pady=5)
        
        search_options = ["Employee ID", "Employee Name", "Department", "Job Title"]
        search_by = Combobox(search_frame, values=search_options, width=15, font=("Helvetica", 11))
        search_by.grid(row=0, column=1, padx=5, pady=5)
        search_by.current(0)  # Default to Employee ID
        
        Label(search_frame, text="Search term:", font=("Helvetica", 11)).grid(row=0, column=2, padx=5, pady=5)
        search_entry = Entry(search_frame, width=20, font=("Helvetica", 11))
        search_entry.grid(row=0, column=3, padx=5, pady=5)
        
        # Create a frame for the treeview
        tree_frame = Frame(data_window)
        tree_frame.pack(fill="both", expand=True, padx=10, pady=5)
        
        # Create scrollbars
        y_scrollbar = Scrollbar(tree_frame, orient="vertical")
        y_scrollbar.pack(side="right", fill="y")
        
        x_scrollbar = Scrollbar(tree_frame, orient="horizontal")
        x_scrollbar.pack(side="bottom", fill="x")
        
        # Create a treeview widget to display the data
        columns = ("رقم الموظف", "الاسم", "القسم", "الدرجة", "الاساسى", 
                   "اجتماعية", "اساسى30/6/15", "اعانه", "بحوث", "ريادة", "اشراف", "مكتبية", 
                   "تطوير", "جودة", "فرق الجودة", "حافز", "بدل", "جملة الاجر", "التاريخ")
        
        tree = Treeview(tree_frame, columns=columns, show="headings",
                            yscrollcommand=y_scrollbar.set,
                            xscrollcommand=x_scrollbar.set)
        
        # Configure scrollbars
        y_scrollbar.config(command=tree.yview)
        x_scrollbar.config(command=tree.xview)
        
        # Format columns
        for col in columns:
            tree.column(col, anchor="w", width=100)
            tree.heading(col, text=col, anchor="w")
        
        tree.pack(fill="both", expand=True)
        
        # Create a frame for buttons
        button_frame = Frame(data_window, padding=10)
        button_frame.pack(fill="x", padx=10, pady=10)
        
        # Function to refresh tree with all data
        def load_all_data():
            # Clear the treeview
            for item in tree.get_children():
                tree.delete(item)
            
            # Fetch all data from the database
            conn = sqlite3.connect('payroll.db')
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM payroll")
            rows = cursor.fetchall()
            conn.close()
            
            # Insert rows into the treeview
            for row in rows:
                tree.insert("", "end", values=row[1:])  # Skip the first column (ID)
        
        # Function to search data
        def search_data():
            search_column = search_by.get()
            search_term = search_entry.get()
            
            if not search_term:
                load_all_data()
                return
            
            # Map UI column names to database column names
            column_mapping = {
                "Employee ID": "employee_id",
                "Employee Name": "employee_name",
                "Department": "department",
                "Job Title": "job_title"
            }
            
            # Get the database column name
            db_column = column_mapping.get(search_column)
            
            if not db_column:
                messagebox.showerror("Error", "Invalid search column selected.")
                return
            
            # Clear the treeview
            for item in tree.get_children():
                tree.delete(item)
            
            # Search in the database
            conn = sqlite3.connect('payroll.db')
            cursor = conn.cursor()
            cursor.execute(f"SELECT * FROM payroll WHERE {db_column} LIKE ?", (f"%{search_term}%",))
            rows = cursor.fetchall()
            conn.close()
            
            # Insert matching rows into the treeview
            for row in rows:
                tree.insert("", "end", values=row[1:])  # Skip the first column (ID)
        
        # Function to export the current view to Excel
        # Function to export the current view to Excel with vertical orientation
        def export_view_to_excel():
            if not tree.get_children():
                messagebox.showinfo("Info", "No data to export.")
                return
            
            try:
                # Ask user where to save the file
                file_path = filedialog.asksaveasfilename(
                    defaultextension=".xlsx",
                    filetypes=[("Excel files", "*.xlsx")],
                    title="Save Excel Report"
                )
                
                if not file_path:
                    return  # User cancelled
                
                # Get the data from the treeview
                data = []
                for item_id in tree.get_children():
                    data.append(tree.item(item_id)['values'])
                
                # Create DataFrame with column names
                df = DataFrame(data, columns=columns)
                
                # Create Excel writer
                with ExcelWriter(file_path, engine='openpyxl') as writer:
                    # Write DataFrame to Excel
                    df.to_excel(writer, index=False, sheet_name="Payroll Data")
                    
                    # Get the workbook and worksheet
                    workbook = writer.book
                    worksheet = writer.sheets["Payroll Data"]
                    
                    # Adjust column widths for better readability
                    for i, col in enumerate(df.columns):
                        column_width = max(df[col].astype(str).map(len).max(), len(col)) + 2
                        worksheet.column_dimensions[openpyxl.utils.get_column_letter(i+1)].width = column_width
                
                messagebox.showinfo("Success", f"Data exported to Excel successfully at {file_path}")
                
            except Exception as e:
                messagebox.showerror("Error", f"Failed to export to Excel: {str(e)}")
        
        # Function to export the current view to PDF with vertical orientation
        # Similarly, update the export_view_to_pdf method in the view_data_from_database function
        def export_view_to_pdf():
            if not tree.get_children():
                messagebox.showinfo("Info", "No data to export.")
                return
        
            try:
                # Ask user where to save the file
                file_path = filedialog.asksaveasfilename(
                    defaultextension=".pdf",
                    filetypes=[("PDF files", "*.pdf")],
                    title="Save PDF Report"
                )
        
                if not file_path:
                    return  # User cancelled
        
                # Import necessary modules
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.ttfonts import TTFont
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
                from reportlab.lib.pagesizes import letter
                from reportlab.lib.units import inch
                from reportlab.lib.colors import grey, whitesmoke, beige
        
                # Register a font that supports Arabic
                arabic_font_path = "Amiri-Regular.ttf"
                pdfmetrics.registerFont(TTFont('Arabic', arabic_font_path))
                pdfmetrics.registerFontFamily('Arabic', normal='Arabic', bold='Arabic', italic='Arabic', boldItalic='Arabic')
                # Create styles
                styles = getSampleStyleSheet()
        
                # Create Arabic paragraph style
                arabic_style = ParagraphStyle(
                    'Arabic',
                    parent=styles['Normal'],
                    fontName='Arabic',
                    alignment=1,  # Center alignment
                    fontSize=10
                )
        
                # Function to format Arabic text
                def format_arabic(text):
                    reshaped_text = arabic_reshaper.reshape(text)  # Reshape the text
                    bidi_text = get_display(reshaped_text)  # Reorder for RTL display
                    return bidi_text
        
                # Create PDF with portrait orientation (default)
                doc = SimpleDocTemplate(file_path, pagesize=letter)
                elements = []
        
                # Add title
                title = Paragraph("Payroll Data Report", styles['Title'])
                elements.append(title)
        
                date_text = Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", 
                                      styles['Normal'])
                elements.append(date_text)
                elements.append(Spacer(1, 12))
        
                # Calculate column widths based on available space
                available_width = letter[0] - inch
                col_width = available_width / 5  # Display 5 columns at a time for readability
        
                # Get the data from the treeview
                data = []
                for item_id in tree.get_children():
                    data.append(tree.item(item_id)['values'])
        
                # Split columns into logical groups for better readability
                column_groups = [
                    columns[0:5],  # Employee info
                    columns[5:10],  # First set of financial data
                    columns[10:15],  # Second set of financial data
                    columns[15:]  # Remaining columns
                ]
        
                group_titles = [
                    "Employee Information",
                    "Financial Data (Part 1)",
                    "Financial Data (Part 2)",
                    "Additional Data"
                ]
        
                # Create multiple tables for different column groups
                for group_idx, column_group in enumerate(column_groups):
                    # Add group title
                    group_title = Paragraph(group_titles[group_idx], styles['Heading2'])
                    elements.append(group_title)
                    elements.append(Spacer(1, 6))
        
                    # Extract data for this column group
                    group_indices = [columns.index(col) for col in column_group]
        
                    # Prepare data for table with proper text handling
                    table_data = []
        
                    # Add header row
                    header_row = []
                    for col in column_group:
                        # Check if column name might contain Arabic
                        if any(ord(c) > 127 for c in col):
                            header_row.append(Paragraph(format_arabic(col), arabic_style))
                        else:
                            header_row.append(Paragraph(col, styles['Normal']))
                    table_data.append(header_row)
        
                    # Add data rows
                    for row in data:
                        table_row = []
                        for i in group_indices:
                            if i < len(row) and row[i] is not None:
                                value = str(row[i])
                                # Check if value might contain Arabic
                                if any(ord(c) > 127 for c in value):
                                    table_row.append(Paragraph(format_arabic(value), arabic_style))
                                else:
                                    table_row.append(Paragraph(value, styles['Normal']))
                            else:
                                table_row.append(Paragraph("", styles['Normal']))
                        table_data.append(table_row)
        
                    # Create table for this group
                    table = Table(table_data, colWidths=[col_width] * len(column_group))
        
                    # Style the table
                    style = TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), whitesmoke),
                        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), beige),
                        ('GRID', (0, 0), (-1, -1), 1, black),
                        ('FONTSIZE', (0, 1), (-1, -1), 8),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE')
                    ])
        
                    table.setStyle(style)
                    elements.append(table)
                    elements.append(Spacer(1, 12))
        
                # Build PDF
                doc.build(elements)
        
                messagebox.showinfo("Success", f"Data exported to PDF successfully at {file_path}")
        
            except Exception as e:
                messagebox.showerror("Error", f"Failed to export to PDF: {str(e)}")
                
        def export_view_to_word():
            if not tree.get_children():
                messagebox.showinfo("Info", "No data to export.")
                return
            
            try:
                # Ask user where to save the file
                file_path = filedialog.asksaveasfilename(
                    defaultextension=".docx",
                    filetypes=[("Word files", "*.docx")],
                    title="Save Word Report"
                )
                
                if not file_path:
                    return  # User cancelled
                
                # Create Word document
                doc = Document()
                doc.add_heading('Payroll Data Report', 0)
                
                # Add current date
                doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                
                # Split columns into logical groups for better readability in portrait mode
                column_groups = [
                    columns[0:5],  # Employee info
                    columns[5:10],  # First set of financial data
                    columns[10:15],  # Second set of financial data
                    columns[15:]  # Remaining columns
                ]
                
                group_titles = [
                    "Employee Information",
                    "Financial Data (Part 1)",
                    "Financial Data (Part 2)",
                    "Additional Data"
                ]
                
                # Get data from the treeview
                data = []
                for item_id in tree.get_children():
                    data.append(tree.item(item_id)['values'])
                
                # Create tables for each column group
                for group_idx, column_group in enumerate(column_groups):
                    # Add section heading
                    doc.add_heading(group_titles[group_idx], level=2)
                    
                    # Create table
                    table = doc.add_table(rows=1, cols=len(column_group))
                    table.style = 'Table Grid'
                    
                    # Add header row
                    header_cells = table.rows[0].cells
                    for i, col in enumerate(column_group):
                        header_cells[i].text = col
                        header_cells[i].paragraphs[0].runs[0].font.bold = True
                    
                    # Add data rows
                    for row_values in data:
                        row_cells = table.add_row().cells
                        for i, col in enumerate(column_group):
                            col_idx = columns.index(col)
                            if col_idx < len(row_values) and row_values[col_idx] is not None:
                                row_cells[i].text = str(row_values[col_idx])
                            else:
                                row_cells[i].text = ""
                    
                    # Add space after table
                    doc.add_paragraph()
                
                # Save document
                doc.save(file_path)
                
                messagebox.showinfo("Success", f"Data exported to Word successfully at {file_path}")
                
            except Exception as e:
                messagebox.showerror("Error", f"Failed to export to Word: {str(e)}")
        
        # Add buttons
        search_button = Button(search_frame, text="Search", command=search_data)
        search_button.grid(row=0, column=4, padx=5, pady=5)
        
        reset_button = Button(search_frame, text="Reset", command=load_all_data)
        reset_button.grid(row=0, column=5, padx=5, pady=5)
        
        # Export buttons
        Button(
            button_frame,
            text="Export to Excel",
            command=export_view_to_excel
        ).pack(side="left", padx=5)
        
        Button(
            button_frame,
            text="Export to PDF",
            command=export_view_to_pdf
        ).pack(side="left", padx=5)
        
        Button(
            button_frame,
            text="Export to Word",
            command=export_view_to_word
        ).pack(side="left", padx=5)
        
        # Load initial data
        load_all_data()



    def create_input_fields(self):
        # Define common departments and job titles for comboboxes
        self.departments = ["جرافيك", "تصوير", "ديكور", "عمارة"]
        self.job_titles = ["أ.د", "أ.م.د", "د", "م.م", "م"]
        
        # Row 0
        Label(self.entry_frame, text="Name:", font=("Helvetica", 12)).grid(row=0, column=0, padx=10, pady=10, sticky="e")
        self.name_entry = Entry(self.entry_frame, width=25, font=("Helvetica", 12), justify="center")
        self.name_entry.grid(row=0, column=1, padx=10, pady=10, sticky="w")
        
        Label(self.entry_frame, text="ID Number:", font=("Helvetica", 12)).grid(row=0, column=2, padx=10, pady=10, sticky="e")
        self.id_entry = Entry(self.entry_frame, width=25, font=("Helvetica", 12), justify="center")
        self.id_entry.grid(row=0, column=3, padx=10, pady=10, sticky="w")
        
        # Row 1
        Label(self.entry_frame, text="Department:", font=("Helvetica", 12)).grid(row=1, column=0, padx=10, pady=10, sticky="e")
        self.department_combo = Combobox(self.entry_frame, values=self.departments, width=22, font=("Helvetica", 12))
        self.department_combo.grid(row=1, column=1, padx=10, pady=10, sticky="w")
        
        Label(self.entry_frame, text="Job Title:", font=("Helvetica", 12)).grid(row=1, column=2, padx=10, pady=10, sticky="e")
        self.job_title_combo = Combobox(self.entry_frame, values=self.job_titles, width=22, font=("Helvetica", 12))
        self.job_title_combo.grid(row=1, column=3, padx=10, pady=10, sticky="w")
        
        # Row 2
        Label(self.entry_frame, text="Basic Salary:", font=("Helvetica", 12)).grid(row=2, column=0, padx=10, pady=10, sticky="e")
        self.basic_salary_entry = Entry(self.entry_frame, width=25, font=("Helvetica", 12), justify="center")
        self.basic_salary_entry.grid(row=2, column=1, padx=10, pady=10, sticky="w")
        
        Label(self.entry_frame, text="Social:", font=("Helvetica", 12)).grid(row=2, column=2, padx=10, pady=10, sticky="e")
        self.social_entry = Entry(self.entry_frame, width=25, font=("Helvetica", 12), justify="center")
        self.social_entry.grid(row=2, column=3, padx=10, pady=10, sticky="w")
        
        # Row 3
        Label(self.entry_frame, text="Basic 30/6/15:", font=("Helvetica", 12)).grid(row=3, column=0, padx=10, pady=10, sticky="e")
        self.basic30_entry = Entry(self.entry_frame, width=25, font=("Helvetica", 12), justify="center")
        self.basic30_entry.grid(row=3, column=1, padx=10, pady=10, sticky="w")
        
        Label(self.entry_frame, text="Enaa:", font=("Helvetica", 12)).grid(row=3, column=2, padx=10, pady=10, sticky="e")
        self.enaa = Entry(self.entry_frame, width=25, font=("Helvetica", 12), justify="center")
        self.enaa.grid(row=3, column=3, padx=10, pady=10, sticky="w")
        

    def init_database(self):
        conn = sqlite3.connect('payroll.db')
        cursor = conn.cursor()
        cursor.execute('''
        CREATE TABLE IF NOT EXISTS payroll (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            employee_id TEXT,
            employee_name TEXT,
            department TEXT,
            job_title TEXT,
            basic_salary REAL,
            Social REAL,
            basic30 REAL,
            enaa REAL,
            bhos REAL,
            ryada REAL,
            eshraf REAL,
            maktabia REAL,
            tatwer REAL,
            gawda REAL,
            diff_gawda REAL,
            hafz REAL,
            badl REAL,
            salary REAL,
            date TEXT
        )
        ''')
        conn.commit()
        conn.close()
    def save_to_database(self):
        if not hasattr(self, 'raw_results'):
            messagebox.showerror("Error", "Calculate payroll first before saving to database.")
            return
        
        try:
            conn = sqlite3.connect('payroll.db')
            cursor = conn.cursor()
            
            # Current date
            current_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            
            # Insert data
            cursor.execute('''
            INSERT INTO payroll (
                employee_id, employee_name, department, job_title, basic_salary, Social, 
                basic30, enaa, bhos, ryada, 
                eshraf, maktabia, tatwer, gawda, diff_gawda, hafz, badl, salary, date
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                self.raw_results["employee_id"],
                self.raw_results["employee_name"],
                self.raw_results["department"],
                self.raw_results["job_title"],
                self.raw_results["basic_salary"],
                self.raw_results["Social"],
                self.raw_results["basic30"],
                self.raw_results["enaa"],
                self.raw_results["bhos"],
                self.raw_results["ryada"],
                self.raw_results["eshraf"],
                self.raw_results["maktabia"],
                self.raw_results["tatwer"],
                self.raw_results["gawda"],
                self.raw_results["diff_gawda"],
                self.raw_results["hafz"],
                self.raw_results["badl"],
                self.raw_results["salary"],
                current_date
            ))
            
            conn.commit()
            conn.close()
            
            messagebox.showinfo("Success", "Payroll data saved to database successfully.")
        except Exception as e:
            messagebox.showerror("Database Error", f"Failed to save to database: {e}")


    def calculate_payroll(self):
        try:
            result = 0
            # Get values from entries and comboboxes
            employee_name = self.name_entry.get()
            employee_id = self.id_entry.get()
            department = self.department_combo.get()
            job_title = self.job_title_combo.get()
            basic_salary = float(self.basic_salary_entry.get())
            social = float(self.social_entry.get())
            basic30 = float(self.basic30_entry.get())
            enaa = float(self.enaa.get())

            # Validate inputs
            if not employee_name or not employee_id:
                messagebox.showerror("Error", "Employee name and ID are required.")
                return
                
            if not department or not job_title:
                messagebox.showerror("Error", "Department and job title are required.")
                return
            # Calculate payroll
            result += (basic30 * 0.775)
            bhos = (basic30 * 0.49)
            result += bhos
            ryada = (basic30 * 0.91)
            result+=ryada
            eshraf = (basic30 * 1.3)
            result+= eshraf
            maktabia = (basic30 * 0.78)
            result += maktabia
            tatwer = (basic30 * 0.78)
            result += tatwer
            gawda = 0
            diff_gawda = 0
            hafz = 0
            badl = 0
            exp = 600
            result += exp
            mnha = 10
            result += mnha
            alawa = 73.90
            result += alawa
            tdress = 1071
            result += tdress
            if job_title == "أ.د":
                badl = 3500
                result+=badl
                gawda = 4270
                diff_gawda = 330
                hafz = 2600
                result += (gawda + diff_gawda + hafz)
            elif job_title == "أ.م.د":
                badl = 3000
                result+=badl
                gawda = 3770
                diff_gawda = 230
                hafz = 2475
                result += (gawda + diff_gawda + hafz)
            elif job_title == "د":
                badl = 2500
                result+=badl
                gawda = 3120
                diff_gawda = 140
                hafz = 2050
                result += (gawda + diff_gawda + hafz)
            elif job_title == "م.م":
                badl = 1500
                result+=badl
                gawda = 2900
                diff_gawda = 100
                hafz = 1850
                result += (gawda + diff_gawda + hafz)
            elif job_title == "م":
                badl = 1000
                result += badl
                gawda = 1850
                diff_gawda = 40
                hafz = 1850
                result += (gawda + diff_gawda + hafz)
            # Store results in dictionary
            self.results = {
                "رقم الموظف": employee_id,
                "الاسم": employee_name,
                "قسم": department,
                "الدرجة الوظيفية": job_title,
                "الاساسى": f"${basic_salary:.2f}",
                "اجتماعية": f"{social}",
                "اساسى 30/6/15": f"{basic30}",
                "اعانة": f"{enaa}%",
                "بحوث": f"${bhos:.2f}",
                "ريادة": f"${ryada:.2f}",
                "اشراف": f"${eshraf:.2f}",
                "مكتبية": f"${maktabia:.2f}",
                "تطوير": f"${tatwer:.2f}",
                "جودة":f"${gawda:.2f}",
                "فرق الجودة":f"${diff_gawda:.2f}",
                "حافز":f"${hafz:.2f}",
                "بدل":f"${badl:.2f}",
                "جملة الاجر":f"${result:.2f}"
            }

            # Save raw values for database
            self.raw_results = {
                "employee_id": employee_id,
                "employee_name": employee_name,
                "department": department,
                "job_title": job_title,
                "basic_salary": basic_salary,
                "Social": social,
                "basic30": basic30,
                "enaa": enaa,
                "bhos":bhos,
                "ryada": ryada,
                "eshraf": eshraf,
                "maktabia": maktabia,
                "tatwer": tatwer,
                "gawda": gawda,
                "diff_gawda": diff_gawda,
                "hafz": hafz,
                "badl":badl,
                "salary":result
            }

            # Display results
            self.display_results()

        except ValueError as e:
            messagebox.showerror("Error", "Please enter valid numeric values for salary, hours, and tax rate.")

    def display_results(self):
        # Clear previous results
        for widget in self.scrollable_frame.winfo_children():
            widget.destroy()
        self.result_widgets = {}

        # Create labels for results
        row = 0
        for key, value in self.results.items():
            label = Label(
                self.scrollable_frame,
                text=f"{key}: ",
                font=("Helvetica", 12, "bold")
            )
            label.grid(row=row, column=0, padx=5, pady=5, sticky="e")
            
            value_label = Label(
                self.scrollable_frame,
                text=value,
                font=("Helvetica", 12)
            )
            value_label.grid(row=row, column=1, padx=5, pady=5, sticky="w")
            
            self.result_widgets[key] = (label, value_label)
            row += 1
            
    # Modify the export_to_pdf method in the PayrollApp class
    def export_to_pdf(self):
        if not hasattr(self, 'results'):
            messagebox.showerror("Error", "Calculate payroll first before exporting.")
            return
    
        try:
            # Ask user where to save the file
            file_path = filedialog.asksaveasfilename(
                defaultextension=".pdf",
                filetypes=[("PDF files", "*.pdf")],
                title="Save PDF Report"
            )
    
            if not file_path:
                return  # User cancelled
    
            # Import necessary modules and dependencies
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.units import inch
            from reportlab.lib.colors import grey, whitesmoke, beige, black
            import arabic_reshaper
            from bidi.algorithm import get_display
            from datetime import datetime
    
            # Register Arabic font and its family mapping so bold/italic are handled properly
            arabic_font_path = "Amiri-Regular.ttf"
            pdfmetrics.registerFont(TTFont('Arabic', arabic_font_path))
            pdfmetrics.registerFontFamily('Arabic', 
                normal='Arabic', 
                bold='Arabic', 
                italic='Arabic', 
                boldItalic='Arabic'
            )
    
            # Create the PDF document
            doc = SimpleDocTemplate(file_path, pagesize=letter)
            styles = getSampleStyleSheet()
    
            # Create an Arabic paragraph style for text that might contain Arabic
            arabic_style = ParagraphStyle(
                'Arabic',
                parent=styles['Normal'],
                fontName='Arabic',
                alignment=1,  # Center alignment
                fontSize=10
            )
    
            elements = []
    
            # Add title and current date to the PDF
            title = Paragraph("Payroll Report", styles['Title'])
            elements.append(title)
            date_text = Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal'])
            elements.append(date_text)
    
            # Prepare data for table output
            data = []
            # Header row
            data.append(["Item", "Value"])
    
            # Function to format Arabic text by reshaping and reordering it for RTL display
            def format_arabic(text):
                reshaped_text = arabic_reshaper.reshape(text)
                bidi_text = get_display(reshaped_text)
                return bidi_text
    
            # Add each key/value from results to the table, applying Arabic formatting if needed
            for key, value in self.results.items():
                if any(ord(c) > 127 for c in key):
                    key_text = Paragraph(format_arabic(key), arabic_style)
                else:
                    key_text = Paragraph(key, styles['Normal'])
    
                if any(ord(c) > 127 for c in value):
                    value_text = Paragraph(format_arabic(value), arabic_style)
                else:
                    value_text = Paragraph(value, styles['Normal'])
    
                data.append([key_text, value_text])
    
            # Create a table with specified column widths
            table = Table(data, colWidths=[3 * inch, 2 * inch])
    
            # Apply styling to the table
            style = TableStyle([
                ('BACKGROUND', (0, 0), (1, 0), grey),
                ('TEXTCOLOR', (0, 0), (1, 0), whitesmoke),
                ('ALIGN', (0, 0), (1, 0), 'CENTER'),
                ('FONTNAME', (0, 0), (1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (1, 0), 14),
                ('BOTTOMPADDING', (0, 0), (1, 0), 12),
                ('BACKGROUND', (0, 1), (1, -1), beige),
                ('GRID', (0, 0), (-1, -1), 1, black),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE')
            ])
            table.setStyle(style)
            elements.append(table)
    
            # Build the PDF
            doc.build(elements)
    
            messagebox.showinfo("Success", f"Payroll report exported to PDF successfully at {file_path}")
    
        except Exception as e:
            messagebox.showerror("Error", f"Failed to export to PDF: {str(e)}")

            
            
    def export_to_excel(self):
        if not hasattr(self, 'results'):
            messagebox.showerror("Error", "Calculate payroll first before exporting.")
            return
    
        try:
            # Ask user where to save the file
            file_path = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Excel files", "*.xlsx")],
                title="Save Excel Report"
            )
            
            if not file_path:
                return  # User cancelled
            
            # Create DataFrame
            data = {"Item": [], "Value": []}
            for key, value in self.results.items():
                data["Item"].append(key)
                data["Value"].append(value)
            
            df = DataFrame(data)
            
            # Write to Excel
            df.to_excel(file_path, index=False, sheet_name="Payroll Report")
            
            messagebox.showinfo("Success", f"Payroll report exported to Excel successfully at {file_path}")
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to export to Excel: {str(e)}")
    def export_to_word(self):
        if not hasattr(self, 'results'):
            messagebox.showerror("Error", "Calculate payroll first before exporting.")
            return
    
        try:
            # Ask user where to save the file
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word files", "*.docx")],
                title="Save Word Report"
            )
            
            if not file_path:
                return  # User cancelled
            
            # Create Word document
            doc = Document()
            doc.add_heading('Payroll Report', 0)
            
            # Add current date
            doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            
            # Add table
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Add header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Item'
            header_cells[1].text = 'Value'
            
            # Add data rows
            for key, value in self.results.items():
                row_cells = table.add_row().cells
                row_cells[0].text = key
                row_cells[1].text = value
            
            # Save document
            doc.save(file_path)
            
            messagebox.showinfo("Success", f"Payroll report exported to Word successfully at {file_path}")
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to export to Word: {str(e)}")

# Main part of the app
if __name__ == "__main__":
    root = Tk()
    app = PayrollApp(root)
    root.mainloop()